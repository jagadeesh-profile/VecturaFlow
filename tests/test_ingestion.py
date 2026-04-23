"""
VecturaFlow — ingestion tests.
Covers: parser (all formats), chunker (metadata preservation, bug fix verification),
lambda_parser handler, edge cases.
No real AWS needed — moto mocks S3, SQS, DynamoDB.
"""
from __future__ import annotations

import csv
import io
import json
import os

from moto import mock_aws
import pytest

# Set env vars before any app import
os.environ.update({
    "OPENAI_API_KEY": "sk-test",
    "PINECONE_API_KEY": "pcsk-test",
    "PINECONE_INDEX": "vecturaflow-test",
    "PINECONE_REGION": "us-east-1",
    "AWS_DEFAULT_REGION": "us-east-1",
    "AWS_ACCESS_KEY_ID": "testing",
    "AWS_SECRET_ACCESS_KEY": "testing",
    "AWS_SECURITY_TOKEN": "testing",
    "AWS_SESSION_TOKEN": "testing",
    "INGESTION_BUCKET": "test-ingestion",
    "EMBEDDING_QUEUE_URL": "https://sqs.us-east-1.amazonaws.com/000000000000/vecturaflow-embedding",
    "REGISTRY_TABLE": "vecturaflow-registry-test",
    "CHUNK_SIZE": "512",
    "CHUNK_OVERLAP": "50",
})

from ingestion.chunker import chunk_blocks, publish_chunks
from ingestion.models import TextBlock
from ingestion.parser import _clean_text, _is_meaningful, parse_file

# ─────────────────────────────────────────────────────────────────────────────
# Helpers to create test file bytes
# ─────────────────────────────────────────────────────────────────────────────

def make_csv_bytes(rows: list[dict]) -> bytes:
    buf = io.StringIO()
    if rows:
        writer = csv.DictWriter(buf, fieldnames=rows[0].keys())
        writer.writeheader()
        writer.writerows(rows)
    return buf.getvalue().encode("utf-8")


def make_txt_bytes(text: str) -> bytes:
    return text.encode("utf-8")


def make_json_bytes(data) -> bytes:
    return json.dumps(data).encode("utf-8")


def make_docx_bytes() -> bytes:
    """Creates a minimal real .docx in memory."""
    from docx import Document
    doc = Document()
    doc.add_heading("VecturaFlow Test Document", level=1)
    doc.add_paragraph("This is the first paragraph about AI and RAG systems.")
    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Retrieval-Augmented Generation combines search with LLMs.")
    doc.add_paragraph("Pinecone stores the vectors. FastAPI serves the results.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# _clean_text tests
# ─────────────────────────────────────────────────────────────────────────────

def test_clean_text_strips_control_chars():
    assert "\x00" not in _clean_text("hello\x00world")
    assert "\x1f" not in _clean_text("test\x1fvalue")


def test_clean_text_normalises_whitespace():
    result = _clean_text("hello    world")
    assert "    " not in result
    assert result == "hello world"


def test_clean_text_collapses_newlines():
    result = _clean_text("para1\n\n\n\n\npara2")
    assert "\n\n\n" not in result


def test_clean_text_empty_input():
    assert _clean_text("") == ""
    assert _clean_text("   ") == ""


def test_is_meaningful_rejects_short():
    assert not _is_meaningful("hi")
    assert not _is_meaningful("123")


def test_is_meaningful_rejects_page_number():
    assert not _is_meaningful("42")
    assert not _is_meaningful("1")


def test_is_meaningful_accepts_real_text():
    assert _is_meaningful("This is a meaningful paragraph about AI systems.")


# ─────────────────────────────────────────────────────────────────────────────
# CSV parser tests
# ─────────────────────────────────────────────────────────────────────────────

def test_parse_csv_basic():
    rows = [
        {"name": "Alice", "role": "Engineer", "team": "AI"},
        {"name": "Bob", "role": "Manager", "team": "Product"},
    ]
    blocks = parse_file(make_csv_bytes(rows), "csv", "doc1", "data.csv")
    assert len(blocks) == 2
    assert all(isinstance(b, TextBlock) for b in blocks)
    assert "Alice" in blocks[0].text
    assert "Engineer" in blocks[0].text
    assert blocks[0].doc_id == "doc1"
    assert blocks[0].source == "data.csv"
    assert blocks[0].file_type == "csv"
    assert blocks[0].row == 0
    assert blocks[1].row == 1


def test_parse_csv_skips_null_rows():
    csv_data = b"name,value\nAlice,Engineer\n,,\nBob,Manager"
    blocks = parse_file(csv_data, "csv", "doc1", "data.csv")
    assert len(blocks) == 2  # empty row skipped


def test_parse_csv_large():
    rows = [{"id": str(i), "text": f"Row {i} content about topic {i}"} for i in range(500)]
    blocks = parse_file(make_csv_bytes(rows), "csv", "doc_large", "large.csv")
    assert len(blocks) == 500


# ─────────────────────────────────────────────────────────────────────────────
# TXT parser tests
# ─────────────────────────────────────────────────────────────────────────────

def test_parse_txt_basic():
    text = "First paragraph about machine learning.\n\nSecond paragraph about deep learning."
    blocks = parse_file(make_txt_bytes(text), "txt", "doc2", "notes.txt")
    assert len(blocks) == 2
    assert "machine learning" in blocks[0].text
    assert "deep learning" in blocks[1].text


def test_parse_txt_single_paragraph():
    text = "This is a long single paragraph. " * 20
    blocks = parse_file(make_txt_bytes(text), "txt", "doc2", "notes.txt")
    assert len(blocks) == 1


def test_parse_txt_filters_noise():
    text = "x\n\n\n\nReal content about VecturaFlow platform.\n\n\n42\n\nMore real content here."
    blocks = parse_file(make_txt_bytes(text), "txt", "doc2", "notes.txt")
    # "x" and "42" should be filtered as too short / pure numeric
    texts = [b.text for b in blocks]
    assert any("VecturaFlow" in t for t in texts)
    assert not any(t.strip() == "x" for t in texts)
    assert not any(t.strip() == "42" for t in texts)


# ─────────────────────────────────────────────────────────────────────────────
# JSON parser tests
# ─────────────────────────────────────────────────────────────────────────────

def test_parse_json_array_of_objects():
    data = [
        {"product": "VecturaFlow", "type": "RAG platform", "status": "active"},
        {"product": "ShaConnects", "type": "Virtual company", "status": "building"},
    ]
    blocks = parse_file(make_json_bytes(data), "json", "doc3", "data.json")
    assert len(blocks) == 2
    assert "VecturaFlow" in blocks[0].text
    assert blocks[0].row == 0
    assert blocks[1].row == 1


def test_parse_json_single_object():
    data = {"name": "Jagadeesh", "role": "AI Engineer", "company": "Seagate"}
    blocks = parse_file(make_json_bytes(data), "json", "doc3", "profile.json")
    assert len(blocks) == 1
    assert "Jagadeesh" in blocks[0].text


def test_parse_json_invalid_raises():
    with pytest.raises(ValueError, match="Invalid JSON"):
        parse_file(b"not valid json {{{", "json", "doc3", "bad.json")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX parser tests
# ─────────────────────────────────────────────────────────────────────────────

def test_parse_docx_basic():
    blocks = parse_file(make_docx_bytes(), "docx", "doc4", "report.docx")
    assert len(blocks) >= 3
    texts = " ".join(b.text for b in blocks)
    assert "VecturaFlow" in texts
    assert "Pinecone" in texts


def test_parse_docx_section_metadata():
    blocks = parse_file(make_docx_bytes(), "docx", "doc4", "report.docx")
    # At least some blocks should have section metadata from headings
    sections = [b.section for b in blocks if b.section]
    assert len(sections) > 0


# ─────────────────────────────────────────────────────────────────────────────
# Unsupported type
# ─────────────────────────────────────────────────────────────────────────────

def test_parse_unsupported_type_raises():
    with pytest.raises(ValueError, match="Unsupported file type"):
        parse_file(b"data", "png", "doc5", "image.png")


# ─────────────────────────────────────────────────────────────────────────────
# ChunkingAgent tests — verifies scorecard bug fix
# ─────────────────────────────────────────────────────────────────────────────

def test_chunk_blocks_preserves_page_metadata():
    """
    CRITICAL: Verifies the scorecard bug fix.
    Each chunk must know which page its text came from.
    """
    blocks = [
        TextBlock(text="A" * 600, doc_id="d1", source="test.pdf", file_type="pdf", page=1),
        TextBlock(text="B" * 600, doc_id="d1", source="test.pdf", file_type="pdf", page=2),
        TextBlock(text="C" * 600, doc_id="d1", source="test.pdf", file_type="pdf", page=3),
    ]
    chunks = chunk_blocks(blocks, chunk_size=512, chunk_overlap=50)

    # All chunks from page-1 block should have page=1
    page1_chunks = [c for c in chunks if "A" in c.text]
    assert all(c.page == 1 for c in page1_chunks), \
        "BUG: page metadata lost — chunks from page 1 should all have page=1"

    page2_chunks = [c for c in chunks if "B" in c.text]
    assert all(c.page == 2 for c in page2_chunks), \
        "BUG: page metadata lost — chunks from page 2 should all have page=2"

    page3_chunks = [c for c in chunks if "C" in c.text]
    assert all(c.page == 3 for c in page3_chunks)


def test_chunk_blocks_global_index_is_unique():
    blocks = [
        TextBlock(text="X" * 600, doc_id="d1", source="f.txt", file_type="txt"),
        TextBlock(text="Y" * 600, doc_id="d1", source="f.txt", file_type="txt"),
    ]
    chunks = chunk_blocks(blocks)
    indices = [c.chunk_index for c in chunks]
    assert indices == list(range(len(chunks))), "chunk_index must be a unique global sequence"


def test_chunk_blocks_total_chunks_backfilled():
    blocks = [TextBlock(text="Z" * 1200, doc_id="d1", source="f.txt", file_type="txt")]
    chunks = chunk_blocks(blocks, chunk_size=512)
    total = len(chunks)
    assert all(c.total_chunks == total for c in chunks), "total_chunks must be backfilled on all chunks"


def test_chunk_blocks_chunk_id_format():
    blocks = [TextBlock(text="Hello world. " * 50, doc_id="abc123", source="f.txt", file_type="txt")]
    chunks = chunk_blocks(blocks)
    for i, chunk in enumerate(chunks):
        assert chunk.chunk_id == f"abc123_chunk_{i}"


def test_chunk_blocks_respects_size():
    blocks = [TextBlock(text="word " * 2000, doc_id="d1", source="f.txt", file_type="txt")]
    chunks = chunk_blocks(blocks, chunk_size=200, chunk_overlap=20)
    for chunk in chunks:
        assert len(chunk.text) <= 250, f"Chunk too large: {len(chunk.text)}"


def test_chunk_blocks_empty_input():
    assert chunk_blocks([]) == []


def test_chunk_blocks_preserves_section():
    blocks = [
        TextBlock(text="Intro content. " * 40, doc_id="d1", source="doc.docx",
                  file_type="docx", section="Introduction"),
        TextBlock(text="Methods content. " * 40, doc_id="d1", source="doc.docx",
                  file_type="docx", section="Methods"),
    ]
    chunks = chunk_blocks(blocks)
    intro_chunks = [c for c in chunks if "Intro" in c.text]
    assert all(c.section == "Introduction" for c in intro_chunks)


@mock_aws
def test_publish_chunks_sends_to_sqs():
    import boto3
    sqs = boto3.client("sqs", region_name="us-east-1")
    queue = sqs.create_queue(QueueName="test-embedding")
    queue_url = queue["QueueUrl"]

    blocks = [TextBlock(text="Test chunk content. " * 30, doc_id="d1", source="f.txt", file_type="txt")]
    chunks = chunk_blocks(blocks)
    published = publish_chunks(chunks, queue_url)

    assert published == len(chunks)

    # Verify messages are in queue
    msgs = sqs.receive_message(QueueUrl=queue_url, MaxNumberOfMessages=10)
    assert len(msgs.get("Messages", [])) == len(chunks)

    # Verify message body has required fields
    msg_body = json.loads(msgs["Messages"][0]["Body"])
    assert "chunk_id" in msg_body
    assert "doc_id" in msg_body
    assert "text" in msg_body
    assert "source" in msg_body
    assert "chunk_index" in msg_body
    assert "total_chunks" in msg_body


# ─────────────────────────────────────────────────────────────────────────────
# Lambda parser handler tests
# ─────────────────────────────────────────────────────────────────────────────

@mock_aws
def test_lambda_parser_processes_csv():
    import importlib

    import boto3
    region = "us-east-1"

    # Setup AWS mocks
    s3 = boto3.client("s3", region_name=region)
    s3.create_bucket(Bucket="test-ingestion")

    dynamo = boto3.resource("dynamodb", region_name=region)
    dynamo.create_table(
        TableName="vecturaflow-registry-test",
        KeySchema=[{"AttributeName": "doc_id", "KeyType": "HASH"}],
        AttributeDefinitions=[{"AttributeName": "doc_id", "AttributeType": "S"}],
        BillingMode="PAY_PER_REQUEST",
    )

    sqs = boto3.client("sqs", region_name=region)
    queue = sqs.create_queue(QueueName="vecturaflow-embedding")
    os.environ["EMBEDDING_QUEUE_URL"] = queue["QueueUrl"]

    # Upload test CSV to S3
    csv_bytes = make_csv_bytes([
        {"name": "Alice", "role": "Engineer", "dept": "AI"},
        {"name": "Bob", "role": "Manager", "dept": "Product"},
        {"name": "Carol", "role": "Analyst", "dept": "Data"},
    ])
    s3.put_object(Bucket="test-ingestion", Key="test/employees.csv", Body=csv_bytes)

    # Reload lambda to pick up mocked clients
    import ingestion.lambda_parser as lp
    importlib.reload(lp)

    event = {"Records": [{
        "messageId": "msg-001",
        "body": json.dumps({
            "doc_id": "testdoc123",
            "bucket": "test-ingestion",
            "key": "test/employees.csv",
            "file_type": "csv",
        })
    }]}

    result = lp.handler(event, None)
    assert result["batchItemFailures"] == []


@mock_aws
def test_lambda_parser_handles_invalid_message():
    import importlib

    import boto3

    dynamo = boto3.resource("dynamodb", region_name="us-east-1")
    dynamo.create_table(
        TableName="vecturaflow-registry-test",
        KeySchema=[{"AttributeName": "doc_id", "KeyType": "HASH"}],
        AttributeDefinitions=[{"AttributeName": "doc_id", "AttributeType": "S"}],
        BillingMode="PAY_PER_REQUEST",
    )

    sqs = boto3.client("sqs", region_name="us-east-1")
    queue = sqs.create_queue(QueueName="vecturaflow-embedding-2")
    os.environ["EMBEDDING_QUEUE_URL"] = queue["QueueUrl"]

    import ingestion.lambda_parser as lp
    importlib.reload(lp)

    # Missing required fields
    event = {"Records": [{
        "messageId": "bad-msg",
        "body": json.dumps({"doc_id": "x"})   # missing bucket, key, file_type
    }]}

    result = lp.handler(event, None)
    assert result["batchItemFailures"] == [{"itemIdentifier": "bad-msg"}]


@mock_aws
def test_lambda_parser_processes_raw_s3_notification():
    import importlib

    import boto3

    region = "us-east-1"

    s3 = boto3.client("s3", region_name=region)
    s3.create_bucket(Bucket="test-ingestion")

    dynamo = boto3.resource("dynamodb", region_name=region)
    dynamo.create_table(
        TableName="vecturaflow-registry-test",
        KeySchema=[{"AttributeName": "doc_id", "KeyType": "HASH"}],
        AttributeDefinitions=[{"AttributeName": "doc_id", "AttributeType": "S"}],
        BillingMode="PAY_PER_REQUEST",
    )

    sqs = boto3.client("sqs", region_name=region)
    queue = sqs.create_queue(QueueName="vecturaflow-embedding-raw")
    os.environ["EMBEDDING_QUEUE_URL"] = queue["QueueUrl"]

    csv_bytes = make_csv_bytes([
        {"name": "Alice", "role": "Engineer", "dept": "AI"},
        {"name": "Bob", "role": "Manager", "dept": "Product"},
    ])
    s3.put_object(Bucket="test-ingestion", Key="incoming/team.csv", Body=csv_bytes)

    import ingestion.lambda_parser as lp
    importlib.reload(lp)

    event = {
        "Records": [{
            "messageId": "s3-msg-001",
            "body": json.dumps({
                "Records": [{
                    "eventSource": "aws:s3",
                    "s3": {
                        "bucket": {"name": "test-ingestion"},
                        "object": {"key": "incoming/team.csv"},
                    },
                }]
            }),
        }]
    }

    result = lp.handler(event, None)
    assert result["batchItemFailures"] == []
